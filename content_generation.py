import os, requests, shutil, json, base64, urllib.parse, warnings
from openpyxl import load_workbook
from datetime import datetime
from docx import Document

#  MICROSOFT API ACCESS TOKEN CODE: This function acquires an access token from Microsoft Graph API using the client credentials flow.
def get_graph_access_token(CLIENT_ID, CLIENT_SECRET,AUTHORITY):
    data = {
      "grant_type": "client_credentials",
      "client_id": CLIENT_ID,
      "scope": "https://graph.microsoft.com/.default",
      "client_secret": CLIENT_SECRET
    }
  
    headers = {"Content-Type": "application/x-www-form-urlencoded"}
    token_url = f"{AUTHORITY}/oauth2/v2.0/token"
    response = requests.post(token_url, headers=headers, data=data)
  
    if response.status_code == 200:
        return response.json()["access_token"]
    else:
        print(f"Error acquiring Graph access token: {response.text}")
        return None

# FILE DOWNLOAD: master function to download a file from SharePoint
def DownloadFile(file_name, access_token, SITE_ID, DOCUMENT_LIBRARY_ID, RESOURCE):
    document_id = search_for_document(access_token, file_name, SITE_ID, DOCUMENT_LIBRARY_ID)
    
    if document_id:
        fetch_sharepoint_document_by_item_id(file_name, access_token, document_id, RESOURCE, SITE_ID, DOCUMENT_LIBRARY_ID)

#  SEARCH FOR DOCUMENT BY NAME: returns the document ID if found, or None if not found.
def search_for_document(access_token, file_name, SITE_ID, DOCUMENT_LIBRARY_ID):
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }

    # URL encode the file name to handle spaces and special characters
    encoded_file_name = urllib.parse.quote(file_name)
    url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/drives/{DOCUMENT_LIBRARY_ID}/root/search(q='{encoded_file_name}')"
    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        items = response.json().get('value', [])
        for item in items:
            if item['name'] == file_name:
                return item['id']
    else:
        return None

# ACCESS SHAREPOINT CODE: This function fetches a document from SharePoint using the item ID.
def fetch_sharepoint_document_by_item_id(file_name, access_token, item_id, RESOURCE, SITE_ID, DOCUMENT_LIBRARY_ID):
    folder_path = 'downloadedFiles'

    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    file_path = os.path.join(folder_path, file_name)

    headers = {
        "Authorization": f"Bearer {access_token}"
    }

    url = f"{RESOURCE}/sites/{SITE_ID}/drives/{DOCUMENT_LIBRARY_ID}/items/{item_id}/content"
    response = requests.get(url, headers=headers, allow_redirects=True)

    if response.status_code == 200:
        with open(file_path, 'wb') as file:
            file.write(response.content)
    else:
        print(f"Error fetching SharePoint document by item ID: {response.text}")

# CHECK FOR AUTHORISED USER FUNCTION: Returns either True or False
def checkAuthorisation(userEmail, userID):
    doc_path = "downloadedFiles/Licensee.xlsx"

    wb = load_workbook(doc_path)
    ws = wb["IDs"]

    id_found = False
    email_matches = False

    for row in ws.iter_rows(min_row=2, values_only=True):
        print("")
        print("Row 7 for ID: " + str(row[7]))
        print("Row 10: '" + str(row[10]) + "' and Row 19: '" + str(row[19]) + "'")

        if userID == row[7]:
            id_found = True
            print("ID found!")
            
            for email in row[19].split('\n'):
                email = email.strip()
                if userEmail == email:
                    print("Email matches!")
                    email_matches = True
                    break 

            if email_matches:
                break


    wb.close()

    if id_found and email_matches:
        return True
    else:
        return False

# CHECK THE LIST FOR WHICH documents HAVE ALREADY BEEN REQUESTED: Returns a list of documents that have been found and a list of documents that have not been found
def checkUnrequestedDocuments(documents_to_search, userID):
    doc_path = "downloadedFiles/Licensee.xlsx"

    # Load the workbook
    wb = load_workbook(doc_path)

    # Load the specific worksheet named 'Content Licensed'
    ws = wb["Content Licensed"]

    id_found = False
    documents_found = []

    # Iterate through rows in the 'Content Licensed' worksheet
    for row in ws.iter_rows(values_only=True):
        if userID in row:
            id_found = True
            for book in documents_to_search:
                if book in row:
                    documents_found.append(book)
                    documents_to_search.remove(book)

    wb.close()

    return documents_found, documents_to_search

# CHECK THE LIST FOR WHICH documents HAVE ALREADY BEEN REQUESTED: Returns a list of documents that have been found and a list of documents that have not been found
def checkUnrequestedDocuments(documents_to_search, userID):
    doc_path = "downloadedFiles/Licensee.xlsx"

    # Load the workbook
    wb = load_workbook(doc_path)

    # Load the specific worksheet named 'Content Licensed'
    ws = wb["Content Licensed"]

    id_found = False
    documents_found = []

    # Iterate through rows in the 'Content Licensed' worksheet
    for row in ws.iter_rows(values_only=True):
        if userID in row:
            id_found = True
            for book in documents_to_search:
                if book in row:
                    documents_found.append(book)
                    documents_to_search.remove(book)

    wb.close()

    return documents_found, documents_to_search

# SAVE AUDIT FUNCTION: Saves the audit of the books that have been requested
def find_first_empty_row(ws, max_check_row=1000):
    for row in range(1, max_check_row + 1):
        if all(ws[row][col].value is None for col in range(ws.max_column)):
            return row
    return ws.max_row + 1  # Default to appending after the last row if no empty row is found

def saveAudit(userID, documents_found, userEmail, userName):
    doc_path = "downloadedFiles/Licensee.xlsx"
    wb = load_workbook(doc_path)
    ws = wb["Content Licensed"]

    start_row = find_first_empty_row(ws)
    current_date = datetime.now().date()
    formatted_date = current_date.strftime('%m/%d/%Y') 

    for i, document in enumerate(documents_found, start=start_row):
        data = [document, userID, userName, userEmail, formatted_date]
        for j, value in enumerate(data, start=1):
            ws.cell(row=i, column=j).value = value

    wb.save(doc_path)
    wb.close()

# FILE DOWNLOAD: master function to download a file from SharePoint
def DownloadFile(file_name, access_token, SITE_ID, DOCUMENT_LIBRARY_ID, RESOURCE):
    document_id = search_for_document(access_token, file_name, SITE_ID, DOCUMENT_LIBRARY_ID)
    
    if document_id:
        fetch_sharepoint_document_by_item_id(file_name, access_token, document_id, SITE_ID, DOCUMENT_LIBRARY_ID, RESOURCE)

#  SEARCH FOR DOCUMENT BY NAME: returns the document ID if found, or None if not found.
def search_for_document(access_token, file_name, SITE_ID, DOCUMENT_LIBRARY_ID):
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }

    # URL encode the file name to handle spaces and special characters
    encoded_file_name = urllib.parse.quote(file_name)
    url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/drives/{DOCUMENT_LIBRARY_ID}/root/search(q='{encoded_file_name}')"
    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        items = response.json().get('value', [])
        for item in items:
            if item['name'] == file_name:
                return item['id']
    else:
        return None

# ACCESS SHAREPOINT CODE: This function fetches a document from SharePoint using the item ID.
def fetch_sharepoint_document_by_item_id(file_name, access_token, item_id, SITE_ID, DOCUMENT_LIBRARY_ID, RESOURCE):
    folder_path = 'downloadedFiles'

    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    file_path = os.path.join(folder_path, file_name)

    headers = {
        "Authorization": f"Bearer {access_token}"
    }

    url = f"{RESOURCE}/sites/{SITE_ID}/drives/{DOCUMENT_LIBRARY_ID}/items/{item_id}/content"
    response = requests.get(url, headers=headers, allow_redirects=True)

    if response.status_code == 200:
        with open(file_path, 'wb') as file:
            file.write(response.content)
    else:
        print(f"Error fetching SharePoint document by item ID: {response.text}")

#  GET DATE TIME FUNCTION: Returns the current date and time
def get_date_with_suffix(date):
    # Get day, month name, and year
    day = date.day
    month_name = date.strftime('%B')
    year = date.year
    # Determine suffix
    if 4 <= day <= 20 or 24 <= day <= 30:
        suffix = "th"
    else:
        suffix = ["st", "nd", "rd"][day % 10 - 1]
    return f"{day}{suffix} {month_name} {year}"

# FETCH BOOK CODE: saves an edited Word File
def WordEditingCode(userID, inputName):
    doc_path = "downloadedFiles/" + inputName
    doc = Document(doc_path)

    current_date = datetime.now()
    formatted_date = get_date_with_suffix(current_date)

    for para in doc.paragraphs:
        if para.text.startswith("Only for use by Licensing Service ID user:"):
            new_text = para.text + " " + userID 
            
            p_runs = len(para.runs)
            for i in range(p_runs):
                para.runs[0].clear()
            
            para.add_run(new_text)

        if "Date of Issue:" in para.text:
            text_to_keep = para.text.split('Date of Issue:')[0] + 'Date of Issue:'
            
            p_runs = len(para.runs)
            for i in range(p_runs):
                para.runs[0].clear()
            
            para.add_run(text_to_keep + " " + formatted_date)


    # Save the document with a new name
    modified_doc_path = "downloadedFiles/" + userID + inputName
    doc.save(modified_doc_path)

# UPLOAD FILE TO SHAREPOINT FUNCTION: Uploads a file to SharePoint
def upload_file_to_sharepoint(access_token, folder_path, local_file_path, SITE_ID, DOCUMENT_LIBRARY_ID):

    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/octet-stream"
    }

    # Extract the file name from the local file path
    file_name = os.path.basename(local_file_path)
    
    # Construct the API URL to upload the file
    upload_url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/drives/{DOCUMENT_LIBRARY_ID}/root:/{folder_path}/{file_name}:/content"

    # Read the file content to be uploaded
    with open(local_file_path, 'rb') as file:
        file_content = file.read()
    
    # Make the PUT request to upload the file
    response = requests.put(upload_url, headers=headers, data=file_content)

    if response.status_code in [200, 201, 202]:
        print(f"Sub - File '{file_name}' uploaded successfully.")
        return response.json()
    else:
        print(f"Error uploading file: {response.status_code}, {response.text}")
        return None